"""
main.py
Orchestrates the DDR pipeline + runs the Flask server.
Includes report writing (docx) and all routes.

Usage:
    python main.py                          # starts Flask server on port 5000
    python main.py --cli --inspection a.pdf --thermal b.pdf --output report.docx
"""

import argparse
import io
import os
import sys
import uuid
from datetime import datetime
from typing import Dict, List

from dotenv import load_dotenv

load_dotenv()


# ── Report writer (python-docx) ────────────────────────────────────────────────

def write_ddr_report(
    ddr_data: dict,
    image_map: Dict[str, list],
    output_path: str,
    property_name: str = "Site Property",
    inspector_name: str = "Not Available",
):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_heading_color(paragraph, r, g, b):
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(r, g, b)

    def add_hr(doc):
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section(doc, title):
        h = doc.add_heading(title, level=2)
        set_heading_color(h, 30, 80, 160)

    def add_image(doc, img, caption=""):
        try:
            stream = io.BytesIO(img.image_bytes)
            doc.add_picture(stream, width=Inches(min(5, img.width / 96)))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                cap = doc.add_paragraph(caption)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.italic = True
                cap.runs[0].font.color.rgb = RGBColor(120, 120, 120)
        except Exception as e:
            doc.add_paragraph(f"[Image could not be rendered: {e}]")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover
    t = doc.add_heading("Detailed Diagnostic Report (DDR)", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_heading_color(t, 20, 60, 140)
    sub = doc.add_paragraph(property_name)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(13)
    meta = doc.add_paragraph(
        f"Report Date: {datetime.today().strftime('%d %B %Y')}     |     Inspector: {inspector_name}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(130, 130, 130)
    add_hr(doc)
    doc.add_paragraph()

    # 1. Summary
    add_section(doc, "1. Property Issue Summary")
    doc.add_paragraph(ddr_data.get("property_issue_summary", "Not Available"))
    doc.add_paragraph()

    # 2. Area-wise Observations
    add_section(doc, "2. Area-wise Observations")
    for area_obs in ddr_data.get("area_wise_observations", []):
        area_name = area_obs.get("area", "Unknown Area")
        h = doc.add_heading(area_name, level=3)
        set_heading_color(h, 60, 120, 60)

        for obs in area_obs.get("observations", []):
            doc.add_paragraph(obs, style="List Bullet")

        area_imgs = image_map.get(area_name, [])
        if area_imgs:
            doc.add_paragraph()
            for img in area_imgs:
                add_image(doc, img, f"Figure: {area_name} — Page {img.page_number}")
                doc.add_paragraph()
        else:
            p = doc.add_paragraph("Image Not Available")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = RGBColor(160, 160, 160)
        doc.add_paragraph()

    # 3. Root Cause
    add_section(doc, "3. Probable Root Cause")
    doc.add_paragraph(ddr_data.get("probable_root_cause", "Not Available"))
    doc.add_paragraph()

    # 4. Severity
    add_section(doc, "4. Severity Assessment")
    severity = ddr_data.get("severity_assessment", {})
    level = severity.get("level", "Not Available")
    color_map = {
        "Low": RGBColor(0, 150, 0),
        "Medium": RGBColor(200, 130, 0),
        "High": RGBColor(200, 60, 0),
        "Critical": RGBColor(180, 0, 0),
    }
    p = doc.add_paragraph()
    run = p.add_run(f"Severity Level: {level}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color_map.get(level, RGBColor(0, 0, 0))
    doc.add_paragraph(f"Reasoning: {severity.get('reasoning', 'Not Available')}")
    doc.add_paragraph()

    # 5. Recommended Actions
    add_section(doc, "5. Recommended Actions")
    actions = ddr_data.get("recommended_actions", [])
    if actions:
        for action in actions:
            doc.add_paragraph(action, style="List Number")
    else:
        doc.add_paragraph("Not Available")
    doc.add_paragraph()

    # 6. Additional Notes
    add_section(doc, "6. Additional Notes")
    doc.add_paragraph(ddr_data.get("additional_notes", "Not Available"))
    doc.add_paragraph()

    # 7. Missing Info
    add_section(doc, "7. Missing or Unclear Information")
    missing = ddr_data.get("missing_or_unclear_information", [])
    if missing:
        for item in missing:
            doc.add_paragraph(item, style="List Bullet")
    else:
        doc.add_paragraph("None — all required information was available.")

    # Appendix: unmatched images
    unmatched = image_map.get("__unmatched__", [])
    if unmatched:
        doc.add_page_break()
        add_section(doc, "Appendix: Additional Images")
        doc.add_paragraph("Images extracted from source documents that could not be assigned to a specific area.")
        doc.add_paragraph()
        for img in unmatched:
            add_image(doc, img, f"Extracted image — Page {img.page_number}")
            doc.add_paragraph()

    add_hr(doc)
    note = doc.add_paragraph(
        "This report was generated by an AI-assisted DDR system. "
        "All findings should be verified by a qualified professional before undertaking remedial works."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(130, 130, 130)

    doc.save(output_path)
    print(f"[INFO] Report saved → {output_path}")


# ── Pipeline ───────────────────────────────────────────────────────────────────

def run_pipeline(inspection_pdf, thermal_pdf, output_path, property_name="Site Property", inspector_name="Not Available"):
    from llm_client import extract_document, call_openrouter, assign_images_to_sections

    print("\n=== DDR Pipeline ===\n")

    print(f"[1/4] Extracting: {inspection_pdf}")
    inspection_doc = extract_document(inspection_pdf)
    print(f"      Pages: {len(inspection_doc.page_texts)} | Images: {len(inspection_doc.images)}")

    print(f"[1/4] Extracting: {thermal_pdf}")
    thermal_doc = extract_document(thermal_pdf)
    print(f"      Pages: {len(thermal_doc.page_texts)} | Images: {len(thermal_doc.images)}")

    print("\n[2/4] Calling LLM...")
    ddr_data = call_openrouter(inspection_doc, thermal_doc)
    print("      Done.")

    print("\n[3/4] Mapping images...")
    image_map = assign_images_to_sections(
        area_observations=ddr_data.get("area_wise_observations", []),
        all_images=inspection_doc.images + thermal_doc.images,
        page_texts_inspection=inspection_doc.page_texts,
        page_texts_thermal=thermal_doc.page_texts,
    )

    print(f"\n[4/4] Writing report → {output_path}")
    write_ddr_report(ddr_data, image_map, output_path, property_name, inspector_name)
    print("\n✅ Done!\n")
    return output_path


# ── Flask server ───────────────────────────────────────────────────────────────

def create_app():
    from flask import Flask, request, jsonify, send_file

    app = Flask(__name__, static_folder="static", static_url_path="")

    UPLOAD_FOLDER = "temp_uploads"
    OUTPUT_FOLDER = "temp_outputs"
    os.makedirs(UPLOAD_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    @app.route("/")
    def index():
        return app.send_static_file("index.html")

    @app.route("/generate", methods=["POST"])
    def generate():
        if "inspection" not in request.files or "thermal" not in request.files:
            return jsonify({"error": "Both inspection and thermal PDF files are required."}), 400

        inspection_file = request.files["inspection"]
        thermal_file = request.files["thermal"]

        if not inspection_file.filename.endswith(".pdf") or not thermal_file.filename.endswith(".pdf"):
            return jsonify({"error": "Only PDF files are accepted."}), 400

        property_name = request.form.get("property_name", "Site Property").strip() or "Site Property"
        inspector_name = request.form.get("inspector_name", "Not Available").strip() or "Not Available"

        job_id = str(uuid.uuid4())[:8]
        inspection_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_inspection.pdf")
        thermal_path = os.path.join(UPLOAD_FOLDER, f"{job_id}_thermal.pdf")
        output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_ddr_report.docx")

        inspection_file.save(inspection_path)
        thermal_file.save(thermal_path)

        try:
            run_pipeline(inspection_path, thermal_path, output_path, property_name, inspector_name)
            return send_file(
                output_path,
                as_attachment=True,
                download_name="DDR_Report.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            return jsonify({"error": str(e)}), 500
        finally:
            for p in [inspection_path, thermal_path]:
                if os.path.exists(p):
                    os.remove(p)

    return app


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--cli", action="store_true", help="Run as CLI instead of server")
    parser.add_argument("--inspection", help="Inspection PDF path (CLI mode)")
    parser.add_argument("--thermal", help="Thermal PDF path (CLI mode)")
    parser.add_argument("--output", default="ddr_report.docx", help="Output path (CLI mode)")
    parser.add_argument("--property", default="Site Property")
    parser.add_argument("--inspector", default="Not Available")
    args = parser.parse_args()

    if not os.environ.get("OPENROUTER_API_KEY"):
        print("[ERROR] OPENROUTER_API_KEY not set in .env")
        sys.exit(1)

    if args.cli:
        if not args.inspection or not args.thermal:
            print("[ERROR] --inspection and --thermal are required in CLI mode.")
            sys.exit(1)
        run_pipeline(args.inspection, args.thermal, args.output, args.property, args.inspector)
    else:
        app = create_app()
        app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
